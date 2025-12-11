"""
Document Processing Service
Handles document upload, text extraction, chunking, and vector storage
"""
from typing import Optional, List, Dict, Any, BinaryIO
from datetime import datetime
import uuid
import os
import tempfile
from pathlib import Path

# Document processing libraries
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd
import openpyxl

from src.config.settings import settings


class DocumentService:
    """Service for document processing and management."""
    
    SUPPORTED_TYPES = {
        "application/pdf": ".pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "application/vnd.ms-excel": ".xls",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
        "text/plain": ".txt",
        "text/csv": ".csv",
    }
    
    def __init__(self, user_id: str):
        self.user_id = user_id
    
    def is_supported(self, content_type: str) -> bool:
        """Check if file type is supported."""
        return content_type in self.SUPPORTED_TYPES
    
    async def process_document(
        self,
        file_content: bytes,
        filename: str,
        content_type: str,
        description: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Process uploaded document:
        1. Extract text
        2. Chunk text
        3. Generate embeddings
        4. Store in Qdrant
        5. Save metadata to PostgreSQL
        """
        if not self.is_supported(content_type):
            raise ValueError(f"Unsupported file type: {content_type}")
        
        document_id = str(uuid.uuid4())
        
        # Extract text based on file type
        text_content = await self._extract_text(
            file_content,
            content_type,
            filename
        )
        
        if not text_content or len(text_content.strip()) < 10:
            raise ValueError("Could not extract meaningful text from document")
        
        # Chunk text for vector storage
        chunks = self._chunk_text(text_content)
        
        # Generate embeddings and store in Qdrant
        await self._store_embeddings(
            document_id=document_id,
            chunks=chunks,
            filename=filename,
        )
        
        # Save document metadata to PostgreSQL
        await self._save_metadata(
            document_id=document_id,
            filename=filename,
            content_type=content_type,
            file_size=len(file_content),
            chunk_count=len(chunks),
            description=description,
        )
        
        # Try to extract financial data
        financial_data = await self._extract_financial_data(text_content)
        
        return {
            "document_id": document_id,
            "filename": filename,
            "content_type": content_type,
            "file_size": len(file_content),
            "chunk_count": len(chunks),
            "text_length": len(text_content),
            "financial_data": financial_data,
            "status": "ready",
        }
    
    async def _extract_text(
        self,
        file_content: bytes,
        content_type: str,
        filename: str,
    ) -> str:
        """Extract text from different file types."""
        
        # PDF
        if content_type == "application/pdf":
            return await self._extract_pdf_text(file_content)
        
        # DOCX
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return await self._extract_docx_text(file_content)
        
        # Excel
        elif content_type in [
            "application/vnd.ms-excel",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ]:
            return await self._extract_excel_text(file_content, filename)
        
        # CSV
        elif content_type == "text/csv":
            return await self._extract_csv_text(file_content)
        
        # Plain text
        elif content_type == "text/plain":
            return file_content.decode("utf-8", errors="ignore")
        
        else:
            raise ValueError(f"Unsupported content type: {content_type}")
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF."""
        import io
        
        text_parts = []
        pdf_file = io.BytesIO(file_content)
        
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            raise ValueError(f"Failed to extract PDF text: {str(e)}")
    
    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX."""
        import io
        
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
            
        except Exception as e:
            raise ValueError(f"Failed to extract DOCX text: {str(e)}")
    
    async def _extract_excel_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from Excel file."""
        import io
        
        try:
            # Try pandas first
            df = pd.read_excel(io.BytesIO(file_content), sheet_name=None)
            
            text_parts = []
            for sheet_name, sheet_df in df.items():
                text_parts.append(f"=== Sheet: {sheet_name} ===")
                text_parts.append(sheet_df.to_string(index=False))
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            raise ValueError(f"Failed to extract Excel text: {str(e)}")
    
    async def _extract_csv_text(self, file_content: bytes) -> str:
        """Extract text from CSV."""
        import io
        
        try:
            df = pd.read_csv(io.BytesIO(file_content))
            return df.to_string(index=False)
            
        except Exception as e:
            raise ValueError(f"Failed to extract CSV text: {str(e)}")
    
    def _chunk_text(
        self,
        text: str,
        chunk_size: int = 1000,
        overlap: int = 200,
    ) -> List[str]:
        """
        Split text into overlapping chunks for vector storage.
        
        Args:
            text: Full text content
            chunk_size: Maximum characters per chunk
            overlap: Overlap between chunks
        
        Returns:
            List of text chunks
        """
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < text_length:
                # Look for sentence ending
                for punct in [". ", ".\n", "! ", "?\n"]:
                    last_punct = text.rfind(punct, start, end)
                    if last_punct > start + chunk_size // 2:
                        end = last_punct + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
        
        return chunks
    
    async def _store_embeddings(
        self,
        document_id: str,
        chunks: List[str],
        filename: str,
    ) -> None:
        """Generate embeddings and store in Qdrant."""
        from src.infrastructure.database.qdrant_client import qdrant_client
        from src.infrastructure.llm.huggingface_client import huggingface_client
        
        # Generate embeddings for all chunks
        embeddings = await huggingface_client.embeddings(chunks)
        
        if not embeddings or len(embeddings) != len(chunks):
            raise ValueError("Failed to generate embeddings")
        
        # Store in Qdrant
        points = []
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            points.append({
                "id": f"{document_id}_{i}",
                "vector": embedding,
                "payload": {
                    "document_id": document_id,
                    "user_id": self.user_id,
                    "filename": filename,
                    "chunk_index": i,
                    "text": chunk,
                    "created_at": datetime.utcnow().isoformat(),
                }
            })
        
        # Batch upsert to Qdrant
        qdrant_client.upsert_batch(points)
    
    async def _save_metadata(
        self,
        document_id: str,
        filename: str,
        content_type: str,
        file_size: int,
        chunk_count: int,
        description: Optional[str] = None,
    ) -> None:
        """Save document metadata to PostgreSQL."""
        from src.infrastructure.database.postgres import async_session_factory
        from src.domain.models.document import Document
        
        async with async_session_factory() as session:
            document = Document(
                id=document_id,
                user_id=self.user_id,
                filename=filename,
                content_type=content_type,
                file_size=file_size,
                chunk_count=chunk_count,
                description=description,
                status="ready",
                created_at=datetime.utcnow(),
            )
            
            session.add(document)
            await session.commit()
    
    async def _extract_financial_data(self, text: str) -> Dict[str, Any]:
        """
        Extract financial data from document text using LLM.
        
        Looks for:
        - Expenses
        - Income
        - Dates
        - Amounts
        - Categories
        """
        import re
        
        financial_data = {
            "expenses": [],
            "income": [],
            "amounts_found": [],
        }
        
        # Find currency amounts
        amount_pattern = r'\$\s*([\d,]+\.?\d*)|(\d+\.?\d*)\s*(?:USD|EUR|GBP|THB)'
        amounts = re.findall(amount_pattern, text)
        
        for match in amounts:
            amount_str = match[0] or match[1]
            try:
                amount = float(amount_str.replace(",", ""))
                if 0 < amount < 1000000:  # Sanity check
                    financial_data["amounts_found"].append(amount)
            except:
                pass
        
        # Use LLM to extract structured data
        if financial_data["amounts_found"]:
            # TODO: Use LLM to extract structured expense/income data
            pass
        
        return financial_data
    
    async def search_documents(
        self,
        query: str,
        top_k: int = 5,
    ) -> List[Dict[str, Any]]:
        """Search user's documents using semantic search."""
        from src.infrastructure.database.qdrant_client import qdrant_client
        from src.infrastructure.llm.huggingface_client import huggingface_client
        
        # Generate query embedding
        embeddings = await huggingface_client.embeddings([query])
        
        if not embeddings or not embeddings[0]:
            raise ValueError("Failed to generate query embedding")
        
        query_vector = embeddings[0]
        
        # Search Qdrant with user filter
        results = qdrant_client.search(
            query_vector=query_vector,
            top_k=top_k,
            filter_conditions={"user_id": self.user_id},
        )
        
        return results
    
    async def get_document(self, document_id: str) -> Optional[Dict[str, Any]]:
        """Get document metadata."""
        from src.infrastructure.database.postgres import async_session_factory
        from src.domain.models.document import Document
        from sqlalchemy import select
        
        async with async_session_factory() as session:
            result = await session.execute(
                select(Document).where(
                    Document.id == document_id,
                    Document.user_id == self.user_id,
                )
            )
            doc = result.scalar_one_or_none()
            
            if not doc:
                return None
            
            return {
                "id": doc.id,
                "filename": doc.filename,
                "content_type": doc.content_type,
                "file_size": doc.file_size,
                "chunk_count": doc.chunk_count,
                "description": doc.description,
                "status": doc.status,
                "created_at": doc.created_at.isoformat(),
            }
    
    async def list_documents(
        self,
        skip: int = 0,
        limit: int = 20,
    ) -> List[Dict[str, Any]]:
        """List user's documents."""
        from src.infrastructure.database.postgres import async_session_factory
        from src.domain.models.document import Document
        from sqlalchemy import select
        
        async with async_session_factory() as session:
            result = await session.execute(
                select(Document)
                .where(Document.user_id == self.user_id)
                .order_by(Document.created_at.desc())
                .offset(skip)
                .limit(limit)
            )
            docs = result.scalars().all()
            
            return [
                {
                    "id": doc.id,
                    "filename": doc.filename,
                    "content_type": doc.content_type,
                    "file_size": doc.file_size,
                    "chunk_count": doc.chunk_count,
                    "description": doc.description,
                    "status": doc.status,
                    "created_at": doc.created_at.isoformat(),
                }
                for doc in docs
            ]
    
    async def delete_document(self, document_id: str) -> bool:
        """Delete document and its embeddings."""
        from src.infrastructure.database.postgres import async_session_factory
        from src.infrastructure.database.qdrant_client import qdrant_client
        from src.domain.models.document import Document
        from sqlalchemy import select, delete
        
        async with async_session_factory() as session:
            # Check document exists and belongs to user
            result = await session.execute(
                select(Document).where(
                    Document.id == document_id,
                    Document.user_id == self.user_id,
                )
            )
            doc = result.scalar_one_or_none()
            
            if not doc:
                return False
            
            # Delete from Qdrant
            qdrant_client.delete_by_filter({"document_id": document_id})
            
            # Delete from PostgreSQL
            await session.execute(
                delete(Document).where(Document.id == document_id)
            )
            await session.commit()
            
            return True
